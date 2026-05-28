from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Inches


SRC = Path("/Users/shixiaolei/Desktop/Project 79 - Financial Information (KE draft 202605xx) - 0521_YL.docx")
OUT = Path("/Users/shixiaolei/Documents/Circleup/project79_work/Project 79 - Financial Information revised markup 20260528.docx")

BLUE = RGBColor(0x00, 0x2F, 0x6C)
RED = RGBColor(0xC0, 0x00, 0x00)


def style_inserted(paragraph, text: str, bold: bool = False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.color.rgb = BLUE
    run.bold = bold
    return paragraph


def append_run(paragraph, text: str, color=BLUE, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return run


def set_visible_markup(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = BLUE


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[-1]
    # The paragraph list is not reliable after XML insertion; wrap the XML node.
    from docx.text.paragraph import Paragraph

    new_para = Paragraph(new_p, paragraph._parent)
    new_para.clear()
    if style:
        new_para.style = style
    if text:
        style_inserted(new_para, text)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return new_para


def insert_table_after(paragraph, rows):
    table = paragraph._parent.add_table(rows=0, cols=len(rows[0]), width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                for run in p.runs:
                    run.font.color.rgb = BLUE
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    return table


def insert_paragraph_after_table(table, text: str = "", style=None):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    new_para = Paragraph(p, table._parent)
    if style:
        new_para.style = style
    if text:
        style_inserted(new_para, text, bold=False)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return new_para


def replace_exact(doc, current: str, new: str):
    current_norm = re.sub(r"\s+", " ", current.strip())
    for p in doc.paragraphs:
        p_norm = re.sub(r"\s+", " ", p.text.strip())
        if p_norm == current_norm:
            style_inserted(p, new)
            return p
    prefix = current_norm[:45]
    for p in doc.paragraphs:
        p_norm = re.sub(r"\s+", " ", p.text.strip())
        if p_norm.startswith(prefix):
            style_inserted(p, new)
            return p
    raise RuntimeError(f"paragraph not found: {current[:80]}")


def find_para(doc, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise RuntimeError(f"paragraph not found: {startswith}")


def replace_text_in_doc(doc):
    replacements = {
        "ur presentation": "Our presentation",
        "set forth our consolidated": "sets forth our consolidated",
        "their accompany disclosures": "the accompanying disclosures",
        "income tax at a preferential tax rate of 15%.": "enterprise income tax at a preferential tax rate of 15% during the validity period of the relevant qualification.",
    }
    agi_defined = False
    for p in doc.paragraphs:
        for run in p.runs:
            t = run.text
            for a, b in replacements.items():
                t = t.replace(a, b)
            if not agi_defined and "AGI" in t:
                t = t.replace("AGI", 'artificial general intelligence ("AGI")', 1)
                agi_defined = True
            t = t.replace("—", "—").replace(" - ", " — ")
            run.text = t
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_after(doc, anchor_start, texts):
    anchor = find_para(doc, anchor_start)
    cur = anchor
    for text in reversed(texts):
        cur = insert_paragraph_after(anchor, text)
    return cur


def main():
    doc = Document(SRC)
    replace_text_in_doc(doc)

    replace_exact(
        doc,
        "[To track finalized OSS]",
        "We are a Specialist Technology Company focusing on AI infrastructure and agentic model-serving technologies. During the Track Record Period, we were at an early commercialisation stage, with revenue beginning in the fourth quarter of 2024 and increasing substantially in 2025 and the three months ended March 31, 2026. We intend to apply for listing as a Pre-Commercial Company under Chapter 18C of the Listing Rules. Accordingly, investors should read this section together with the disclosures on our Specialist Technology Products, research and development activities, commercialisation progress, working capital, cash operating costs and R&D expenditure set out in this document.",
    )

    add_after(
        doc,
        "Our results of operations, financial condition",
        [
            "Commercialisation Stage and Path to Profitability",
            "We commenced commercialisation in the fourth quarter of 2024 and continued to scale our revenue-generating business in 2025 and the three months ended March 31, 2026. Our revenue growth during the Track Record Period was primarily driven by (i) the transition from early-stage commercialisation to a broader customer base for our platform software and AGI enterprise solutions, (ii) the launch and rapid ramp-up of Agentic MaaS, which is charged primarily based on customer usage, and (iii) increasing demand from large model companies and enterprise customers for model training, inference and deployment capabilities. We expect our path to profitability to depend principally on continued revenue growth, higher utilisation of computing resources, optimisation of computing resource procurement and leasing arrangements, improved gross profit margin of MaaS and solution projects, and disciplined control of personnel, sales and administrative expenses. We expect to continue to incur substantial R&D expenses in the near term as we invest in our core technologies and products.",
        ],
    )

    replace_exact(
        doc,
        "Our cost of sales primarily consisted of [●]. During the Track Record Period, our cost of sales and services was nil, RMB2.2 million, RMB33.9 million, RMB[●] million and RMB64.7 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively, primarily driven by our revenue growth. The following table sets forth a breakdown of our cost of sales by nature, both in absolute amount and as a percentage of total cost of sales, for the periods indicated.",
        "Our cost of sales primarily consisted of computing resource procurement and leasing costs, server and cloud service costs, project delivery costs and other costs directly attributable to the delivery of our products and services. During the Track Record Period, our cost of sales and services was nil, RMB2.2 million, RMB33.9 million, RMB[●] million and RMB64.7 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively. The increase was primarily driven by our revenue growth and the change in business mix, particularly the rapid growth of Agentic MaaS and other businesses that require computing resource inputs. Platform software generally involved relatively limited direct costs, while MaaS and computing-resource-intensive solution projects incurred higher direct costs, primarily because they require access to leased or procured computing resources. The following table sets forth a breakdown of our cost of sales by nature, both in absolute amount and as a percentage of total cost of sales, for the periods indicated.",
    )

    replace_exact(
        doc,
        "[to supplement analysis of significantly high/low gross profit margin]",
        "Our gross profit margin decreased from 77.7% in 2024 to 55.2% in 2025 and further to 20.6% for the three months ended March 31, 2026, primarily due to changes in our business mix and commercialisation stage. Platform software typically generated a higher gross profit margin because it required limited incremental direct costs after product development. By contrast, Agentic MaaS and certain solution projects required computing resource leasing or procurement, server resources and related delivery inputs. As Agentic MaaS became a more meaningful revenue contributor in 2025 and grew significantly in the first quarter of 2026, the relative contribution of computing-resource-intensive services increased, resulting in a lower overall gross profit margin. Management expects gross profit margin to improve over time if we achieve higher resource utilisation, optimise computing resource procurement and leasing arrangements, improve scheduling efficiency and benefit from scale effects, although the timing and extent of improvement will depend on customer demand, pricing, resource costs and business mix.",
    )

    replace_exact(
        doc,
        "Our selling and distribution expenses primarily consisted of [●]. We incurred selling and distribution of RMB2.7 million, RMB32.9 million, RMB84.3 million, RMB[●] million and RMB13.0 million in 2023, 2024, 2025 and for the three months ended March 31, 2025 and 2026, respectively. The following table sets forth a breakdown of our selling and distribution expenses, both in absolute amount and as a percentage of total selling and distribution expenses, for the periods indicated.",
        "Our selling and distribution expenses primarily consisted of employee benefit expenses for sales and business development personnel, market promotion expenses, customer acquisition and business development expenses and other commercialisation-related expenses. We incurred selling and distribution expenses of RMB2.7 million, RMB32.9 million, RMB84.3 million, RMB[●] million and RMB13.0 million in 2023, 2024, 2025 and for the three months ended March 31, 2025 and 2026, respectively, representing [●]%, 325.6%, 111.6%, [●]% and 16.0% of our revenue for the respective periods. The increase during the Track Record Period was primarily attributable to the expansion of our sales and business development team, increased market promotion and customer engagement activities and higher commercialisation spending to support business growth. The following table sets forth a breakdown of our selling and distribution expenses, both in absolute amount and as a percentage of total selling and distribution expenses, for the periods indicated.",
    )

    replace_exact(
        doc,
        "Our administrative expenses primarily consisted of [●]. We incurred administrative expenses of RMB14.7 million, RMB69.3 million, RMB82.9 million, RMB[●] million and RMB20.3 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively. The following table sets forth a breakdown of our administrative expenses, both in absolute amount and as a percentage of total administrative expenses, for the periods indicated",
        "Our administrative expenses primarily consisted of employee benefit expenses for administrative and management personnel, professional service fees, office and general administrative expenses, depreciation and amortisation and other compliance and corporate operation expenses. We incurred administrative expenses of RMB14.7 million, RMB69.3 million, RMB82.9 million, RMB[●] million and RMB20.3 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively, representing [●]%, 687.4%, 109.7%, [●]% and 25.0% of our revenue for the respective periods. The increase from 2023 to 2024 was primarily due to the expansion of our personnel and the need to engage third-party professional advisers to support our business growth and compliance functions. The further increase in 2025 primarily reflected higher personnel costs and continued business expansion, while administrative expenses as a percentage of revenue decreased as revenue scaled. The following table sets forth a breakdown of our administrative expenses, both in absolute amount and as a percentage of total administrative expenses, for the periods indicated.",
    )

    replace_exact(
        doc,
        "Our research and development expenses primarily consisted of [●]. We incurred research and development expenses of RMB30.6 million, RMB252.9 million, RMB354.7 million, RMB[●] million and RMB93.4 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively. The following table sets forth a breakdown of our research and development expenses, both in absolute amount and as a percentage of total research and development expenses, for the periods indicated.",
        "Our research and development expenses primarily consisted of employee benefit expenses for R&D personnel, server and computing resource costs, technical service fees, cloud service fees, depreciation and amortisation, share-based payment expenses and other expenses incurred for our core technology and product development activities. We incurred research and development expenses of RMB30.6 million, RMB252.9 million, RMB354.7 million, RMB[●] million and RMB93.4 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively, representing [●]%, 2,509.2%, 469.5%, [●]% and 114.6% of our revenue for the respective periods. The significant increase during the Track Record Period was primarily due to the expansion of our R&D team, continued investment in heterogeneous computing scheduling, model training and inference optimisation, MaaS platform capabilities, software and hardware coordination and other frontier technologies, and certain major R&D projects. We did not capitalise our R&D expenses during the Track Record Period. The following table sets forth a breakdown of our research and development expenses, both in absolute amount and as a percentage of total research and development expenses, for the periods indicated.",
    )

    replace_exact(
        doc,
        "Our research and development represented 2,509.2%, 469.5%, [●]% and 114.6% of our revenue in 2024, 2025 and for the three months ended March 31, 2025 and 2026, respectively. [To discuss reasons]",
        "Our research and development expenses represented 2,509.2%, 469.5%, [●]% and 114.6% of our revenue in 2024, 2025 and for the three months ended March 31, 2025 and 2026, respectively. The ratio decreased as revenue scaled from a low base, although R&D expenses remained substantial in absolute amount. We expect to continue to invest in R&D to support product iteration and commercialisation, while seeking to improve operating leverage as revenue grows.",
    )

    replace_exact(
        doc,
        "Our revenue increased significantly from RMB10.1 million in 2024 to RMB75.5 million in 2025 due to the expansion of our agentic platform software and AGI enterprise solutions business. We also started to generate revenue from the provision of Agentic MaaS in 2025.",
        "Our revenue increased significantly from RMB10.1 million in 2024 to RMB75.5 million in 2025, primarily because 2024 represented the early stage of our commercialisation and we began broader commercialisation only in the fourth quarter of 2024, while 2025 was the first full year in which we commercialised our products and services. The increase was driven by the expansion of our agentic platform software and AGI enterprise solutions business, growth in customer number and project opportunities, and the commencement of revenue contribution from Agentic MaaS in 2025.",
    )

    replace_exact(
        doc,
        "Our revenue increased significantly from nil in 2023 to RMB10.1 million in 2024 since we started to generate revenue from provision of agentic platform software and AGI enterprise solutions.",
        "Our revenue increased significantly from nil in 2023 to RMB10.1 million in 2024 as we commenced commercialisation and started to generate revenue from the provision of agentic platform software and AGI enterprise solutions in the fourth quarter of 2024.",
    )

    replace_exact(
        doc,
        "Our property, plant and equipment primarily consisted of electronic equipment and office furniture. The following table sets forth the components of our property, plant and equipment as of the dates indicated.",
        "Our property, plant and equipment primarily consisted of electronic equipment and office furniture. Our electronic equipment mainly included GPU servers, CPU servers, storage equipment, network security equipment and other equipment used to support our R&D activities and internal computing environment. The following table sets forth the components of our property, plant and equipment as of the dates indicated.",
    )

    replace_exact(
        doc,
        "Our right-of-use assets consisted of leased properties and server device. Our right-of-use assets increased from RMB5.6 million as of December 31, 2023 to RMB538.6 million as of December 31, 2024, primarily due to purchase of additional server device, in line with [●]. Our right-of-use assets decreased to RMB446.7 million as of December 31, 2025, primarily due to [depreciation of server device]. Our right-of-use assets increased to RMB 472.3 million as of March 31, 2026, since [●].",
        "Our right-of-use assets consisted of leased properties and server devices. Our right-of-use assets increased from RMB5.6 million as of December 31, 2023 to RMB538.6 million as of December 31, 2024, primarily because we entered into server and computing resource leasing arrangements with lease terms exceeding one year to support our R&D activities, model training and inference capabilities, MaaS services and AI infrastructure deployment. Our right-of-use assets decreased to RMB446.7 million as of December 31, 2025, primarily due to depreciation of server devices. Our right-of-use assets increased to RMB472.3 million as of March 31, 2026, primarily due to the renewal of office premises leases and other lease arrangements. We generally use leasing arrangements for certain server and computing resources because leasing allows us to spread cash outflows over time and provides greater flexibility than direct purchase of relevant assets.",
    )

    replace_exact(
        doc,
        "Our intangible assets consisted of software and patent. Our intangible assets increased from RMB0.01 million as of December 31, 2023 to RMB0.1 million as of December 31, 2024, and further to RMB115.4 million as of December 31, 2025, primarily attributable to [the recognition of intellectual property rights contributed to us in 2025]. Our intangible assets remained relatively stable at RMB111.2 million as of March 31, 2026.",
        "Our intangible assets consisted of software and patents. Our intangible assets increased from RMB0.01 million as of December 31, 2023 to RMB0.1 million as of December 31, 2024, and further to RMB115.4 million as of December 31, 2025, primarily attributable to the recognition of intellectual property rights contributed to us by Tsinghua University and Shanghai Jiao Tong University in 2025, which were recorded based on the valuation agreed by the relevant parties and are amortised over ten years. Such intellectual property rights are related to our core R&D technologies. Our intangible assets remained relatively stable at RMB111.2 million as of March 31, 2026.",
    )

    replace_exact(
        doc,
        "Our trade and other receivables increased from RMB50.3 million as of December 31, 2023 to RMB83.9 million as of December 31, 2024, further to RMB187.2 million and RMB337.8 million as of December 31, 2025 and March 31, 2026, in line with our business expansion.",
        "Our trade and other receivables increased from RMB50.3 million as of December 31, 2023 to RMB83.9 million as of December 31, 2024, and further to RMB187.2 million and RMB337.8 million as of December 31, 2025 and March 31, 2026, respectively, primarily in line with our business expansion and revenue growth. The increase as of March 31, 2026 also reflected the rapid growth of revenue in the first quarter of 2026, certain trade receivables that remained within the credit period, and increases in other receivables and deposits relating to computing resource arrangements under which certain resources were procured or leased on a prepayment or deposit basis due to market supply conditions.",
    )

    replace_exact(
        doc,
        "Our cash and cash equivalents consisted of short-term bank deposits. Our cash and cash equivalents increased from RMB147.6 million as of December 31, 2023 to RMB354.9 million as of December 31, 2024, and further to RMB566.2 million as of December 31, 2025, primarily due to [issuance of shares with preferential rights and bank borrowings]. Our cash and cash equivalents decreased to RMB529.2 million as of March 31, 2026, since [●].",
        "Our cash and cash equivalents consisted of short-term bank deposits. Our cash and cash equivalents increased from RMB147.6 million as of December 31, 2023 to RMB354.9 million as of December 31, 2024, and further to RMB566.2 million as of December 31, 2025, primarily due to financing activities, including issuance of shares with preferential rights and bank borrowings. Our cash and cash equivalents decreased to RMB529.2 million as of March 31, 2026, primarily due to cash used in our daily operations, R&D activities and commercialisation activities. As of May 21, 2026, we had completed the closing of RMB500 million of Series B financing proceeds, and we had obtained domestic bank credit facilities of approximately RMB500 million. These funding sources support our liquidity and working capital needs. [Company and Reporting Accountants to confirm amount, timing and disclosure presentation.]",
    )

    replace_exact(
        doc,
        "During the Track Record Period, our primary use of cash was to [fund our R&D, procurement of computing resources, sales and business development activities, employee costs and other operational needs]. We financed our operations and other capital requirements mainly through [cash generated from our business operations, bank borrowings and equity financing]. We do not anticipate any changes to the availability of financing to fund our operations in the future.",
        "During the Track Record Period, our primary use of cash was to fund our R&D activities, employee benefit expenses, procurement and leasing of computing resources, sales and business development activities, lease payments, capital expenditures and other operational needs. We financed our operations and other capital requirements mainly through equity financing, bank borrowings and, to a lesser extent, cash generated from revenue-generating activities. We will continue to manage our cash outflows with reference to our commercialisation progress, computing resource requirements and R&D plan.",
    )

    add_after(
        doc,
        "Our anticipated cash needs primarily relate",
        [
            "Our current cash burn primarily consists of employee compensation, R&D investment, server and computing resource leasing or procurement, cloud and technical service costs, lease payments, capital expenditures for R&D equipment and other daily operating expenses. We expect our cash runway to be supported by our cash and cash equivalents, Series B financing proceeds received after the Track Record Period, available banking facilities, revenue generated from our products and services and the estimated net proceeds from the Global Offering. [Company and Reporting Accountants to confirm cash burn amount and runway period.]",
        ],
    )

    # Add Chapter 18C-style disclosure sections after cash operating activity discussion anchor.
    anchor = find_para(doc, "Net Cash Flows [Used in] Operating Activities")
    cash_h = insert_paragraph_after(anchor, "Cash Operating Costs", style=anchor.style)
    cash_p = insert_paragraph_after(
        cash_h,
        "The table below sets forth key information relating to our cash operating costs for the periods indicated. Cash operating costs are intended to present the principal cash outflows used to support our operations and commercialisation, including R&D, workforce employment, computing resource costs and other significant operating costs. The calculation below should be reconciled to the Accountants' Report and cash flow statements. [Reporting Accountants to confirm amounts and adjustments.]",
    )
    rows = [
        ["", "2023", "2024", "2025", "Three months ended March 31, 2026", "Basis / note"],
        ["Research and development cash costs", "[●]", "[●]", "[●]", "[●]", "R&D expenses excluding non-cash items and employee compensation, adjusted for related prepayments and payables."],
        ["Workforce employment", "[●]", "[●]", "[●]", "[●]", "Employee benefit expenses under R&D, selling, administrative and cost of sales, excluding non-cash share-based payments, adjusted for staff cost payables."],
        ["Computing resource and direct delivery costs", "[●]", "[●]", "[●]", "[●]", "Cost of sales and direct service costs relating to server, cloud, computing resource leasing/procurement and project delivery, excluding non-cash items and adjusted for related prepayments and payables."],
        ["Sales, marketing and commercialisation", "[●]", "[●]", "[●]", "[●]", "Selling and distribution expenses excluding employee compensation and non-cash items, adjusted for related payables."],
        ["Other significant operating costs", "[●]", "[●]", "[●]", "[●]", "Office, compliance, professional service, lease and other significant operating cash costs not included above."],
        ["Total cash operating costs", "[●]", "[●]", "[●]", "[●]", ""],
    ]
    cash_table = insert_table_after(cash_p, rows)

    rd_h = insert_paragraph_after_table(cash_table, "R&D Expenditure and Total Operating Expenditure", style=anchor.style)
    rd_intro = insert_paragraph_after(
        rd_h,
        "For purposes of Chapter 18C of the Listing Rules, a Pre-Commercial Company is expected to demonstrate substantial R&D investment during the track record period. Based on the current financial statement line items and subject to confirmation by the Reporting Accountants, our R&D expenditure ratio calculated using research and development expenses divided by the sum of research and development expenses, selling and distribution expenses and administrative expenses was approximately 63.7%, 71.2%, 68.0% and 73.7% in 2023, 2024, 2025 and for the three months ended March 31, 2026, respectively, and approximately 69.0% for 2023 to 2025 in aggregate. The table below should be updated if the Reporting Accountants determine that any adjustments, including capitalised development expenditure, non-operating items, listing expenses or other items, should be included or excluded for Chapter 18C purposes.",
    )
    rows2 = [
        ["", "2023", "2024", "2025", "Three months ended March 31, 2026", "2023-2025 total"],
        ["Research and development expenses / R&D expenditure", "30,596", "252,924", "354,653", "93,361", "638,173"],
        ["Selling and distribution expenses", "2,700", "32,900", "84,300", "13,000", "119,900"],
        ["Administrative expenses", "14,700", "69,300", "82,900", "20,300", "166,900"],
        ["Total operating expenditure", "47,996", "355,124", "521,853", "126,661", "924,973"],
        ["R&D expenditure ratio", "63.7%", "71.2%", "68.0%", "73.7%", "69.0%"],
    ]
    insert_table_after(rd_intro, rows2)

    # Mark new/replaced paragraphs in blue for visibility.
    for p in doc.paragraphs:
        if any(run.font.color.rgb == BLUE for run in p.runs):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
