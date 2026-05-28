from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

BLUE = RGBColor(0x00, 0x2F, 0x6C)

SRC = Path("/Users/shixiaolei/Documents/Circleup/project79_work/Project 79 - Financial Information revised markup 20260528.docx")
OUT = Path("/Users/shixiaolei/Documents/Circleup/project79_work/Project 79 - Financial Information revised markup with AR 20260528.docx")


def mark_para(p, text):
    p.clear()
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def insert_after(p, text, style=None):
    from copy import deepcopy
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(p._p)
    p._p.addnext(new_p)
    para = Paragraph(new_p, p._parent)
    para.clear()
    if style is not None:
        para.style = style
    mark_para(para, text)
    return para


def find(doc, starts):
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts):
            return p
    raise RuntimeError(starts)


def replace_start(doc, starts, text):
    p = find(doc, starts)
    mark_para(p, text)
    return p


def replace_contains(doc, needle, text):
    for p in doc.paragraphs:
        if needle in p.text:
            mark_para(p, text)
            return p
    raise RuntimeError(needle)


def update_table_cells(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells and cells[0] == "Inventories":
                cells[0].paragraphs[0].clear()
                r = cells[0].paragraphs[0].add_run("Inventories / contract costs")
                r.font.color.rgb = BLUE
            if cells and cells[0] == "Contract costs" and len(cells) >= 6:
                # Accountants' report supports 2023-2025; leave 2026/latest placeholders unchanged.
                vals = ["Contract costs", "—", "—", "—", "785", "1,970"]
                for i, v in enumerate(vals[: len(row.cells)]):
                    row.cells[i].text = v
                    for p in row.cells[i].paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = BLUE
            if cells and "Interest on financial liabilities on shares" in cells[0] and len(cells) >= 4:
                vals = ["Interest on financial liabilities on shares with preferential rights", "9,904", "64,055", "111,282"]
                for i, v in enumerate(vals[: len(row.cells)]):
                    row.cells[i].text = v
                    for p in row.cells[i].paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = BLUE


def main():
    doc = Document(SRC)

    replace_start(
        doc,
        "The revenue from our agentic platform software is recognized",
        "Revenue from our agentic platform software is recognised at a point in time when the software is delivered to the customer, inspected and accepted by the customer. According to the Accountants' Report, payment is generally due within one to three months from delivery.",
    )
    replace_start(
        doc,
        "The revenue from our AGI enterprise solutions is recognized",
        "Revenue from our AGI enterprise solutions is recognised at a point in time when the software platform and related services are delivered to the customer's designated place, inspected and accepted by the customer. According to the Accountants' Report, payment is generally due within one to three months from delivery.",
    )
    replace_start(
        doc,
        "The revenue from our Agentic MaaS is recognized",
        "Revenue from our Agentic MaaS is recognised as services are rendered based on the customer's consumption of resources for usage-based contracts.",
    )

    replace_start(
        doc,
        "Our other income primarily consisted",
        "Our other income primarily consisted of (i) government grants, (ii) interest income and (iii) others. According to the Accountants' Report, our government grants amounted to RMB1.2 million, RMB22.0 million and RMB23.2 million in 2023, 2024 and 2025, respectively, and mainly represented subsidies received from local governments in Beijing and Shanghai to reward contributions by our subsidiaries to local economic growth or for industry development. The Accountants' Report states that there were no unfulfilled conditions or contingencies relating to these grants. Our interest income amounted to nil, RMB7.5 million and RMB8.3 million in 2023, 2024 and 2025, respectively. Others amounted to nil, RMB7.8 million and RMB4.2 million in 2023, 2024 and 2025, respectively. The following table sets forth the breakdown of our other income in absolute amount and as a percentage of total other income for the periods indicated.",
    )

    replace_start(
        doc,
        "Our finance costs primarily represented",
        "Our finance costs primarily represented interest on financial liabilities on shares with preferential rights, lease liabilities and bank borrowings. We incurred finance costs of RMB9.9 million, RMB70.3 million, RMB134.1 million, RMB[●] million and RMB44.9 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively. According to the Accountants' Report, interest on financial liabilities on shares with preferential rights was RMB9.9 million, RMB64.1 million and RMB111.3 million in 2023, 2024 and 2025, respectively; interest on lease liabilities was nil, RMB6.3 million and RMB21.2 million in 2023, 2024 and 2025, respectively; and interest on bank borrowings was nil, RMB25 thousand and RMB1.6 million in 2023, 2024 and 2025, respectively. The following table sets forth a breakdown of our finance costs for the periods indicated.",
    )

    replace_start(
        doc,
        "Under the Law of the PRC on Enterprise Income Tax",
        "Under the Law of the PRC on Enterprise Income Tax (\"EIT Law\") and Implementation Regulation of the EIT Law, the tax rate of the PRC subsidiaries is generally 25% during the Track Record Period, except for entities entitled to preferential tax treatment. According to the Accountants' Report, Shanghai Infinigence was qualified as a high and new technology enterprise (\"HNTE\") on December 25, 2025 and Beijing Wuwen was qualified as an HNTE on December 30, 2025, and each was entitled to a preferential tax rate of 15% from 2025 to 2027. Beijing Wuwen and Shanghai Infinigence were also entitled to the R&D expense super deduction during the Track Record Period.",
    )

    replace_start(
        doc,
        "Our property, plant and equipment primarily consisted",
        "Our property, plant and equipment primarily consisted of electronic equipment, office furniture and others. Our electronic equipment mainly included GPU servers, CPU servers, storage equipment, network security equipment and other equipment used to support our R&D activities and internal computing environment. According to the Accountants' Report, our electronic equipment is depreciated at 33.33% per annum, office furniture at 20% to 25% per annum and others at 33.33% per annum on a straight-line basis after taking into account residual values. The following table sets forth the components of our property, plant and equipment as of the dates indicated.",
    )

    replace_start(
        doc,
        "Our right-of-use assets consisted",
        "Our right-of-use assets consisted of leased properties and server devices. According to the Accountants' Report, our right-of-use assets increased from RMB5.6 million as of December 31, 2023 to RMB538.6 million as of December 31, 2024, primarily reflecting additions of RMB591.4 million in server devices in 2024, and decreased to RMB446.7 million as of December 31, 2025, primarily reflecting depreciation charges, including RMB146.5 million for server devices in 2025, partially offset by additions. Lease contracts are generally entered into for fixed terms of one to five years and may contain extension and termination options. We generally use leasing arrangements for certain server and computing resources because leasing allows us to spread cash outflows over time and provides greater flexibility than direct purchase of relevant assets.",
    )

    replace_start(
        doc,
        "Our intangible assets consisted",
        "Our intangible assets consisted of software and patents. According to the Accountants' Report, our intangible assets increased from RMB0.01 million as of December 31, 2023 to RMB0.1 million as of December 31, 2024, and further to RMB115.4 million as of December 31, 2025, primarily due to additions of RMB12.5 million in software and RMB110.6 million in patents in 2025, partially offset by amortisation. The Accountants' Report states that these intangible assets have finite useful lives and are amortised on a straight-line basis over two to ten years. As advised by the Company, the patent additions were related to intellectual property rights contributed by Tsinghua University and Shanghai Jiao Tong University and are related to our core R&D technologies.",
    )

    replace_start(
        doc,
        "Our inventories consisted primarily",
        "According to the Accountants' Report, we did not record inventories as of December 31, 2023, 2024 or 2025. We recorded contract costs of nil, RMB0.8 million and RMB2.0 million as of December 31, 2023, 2024 and 2025, respectively, representing costs to fulfil contracts, which mainly comprised engineering costs incurred directly related to existing contracts that will be used to satisfy performance obligations in the future. Based on management's FDDQ response, the RMB3.1 million balance as of March 31, 2026 described as inventories in the management accounts represented costs incurred for solution projects before the relevant revenue recognition milestone was reached and would be transferred to cost of sales when the project reaches the revenue recognition point. [Reporting Accountants to confirm the March 31, 2026 classification.]",
    )
    replace_contains(
        doc,
        "Our inventories [increased]",
        "Our contract costs increased from nil as of December 31, 2023 to RMB0.8 million as of December 31, 2024 and further to RMB2.0 million as of December 31, 2025, primarily due to the increase in engineering costs incurred for solution projects that had not yet reached the relevant revenue recognition point. As of March 31, 2026, the corresponding balance was RMB3.1 million based on management accounts, reflecting the continued execution of solution projects before acceptance or other revenue recognition milestones. [Reporting Accountants to confirm classification and balance.]",
    )

    replace_start(
        doc,
        "Our trade and other receivables increased",
        "Our trade and other receivables increased from RMB50.3 million as of December 31, 2023 to RMB83.9 million as of December 31, 2024, and further to RMB187.2 million and RMB337.8 million as of December 31, 2025 and March 31, 2026, respectively, primarily in line with our business expansion and revenue growth. According to the Accountants' Report, trade receivables net of allowances increased from nil as of December 31, 2023 to RMB0.6 million as of December 31, 2024 and RMB41.4 million as of December 31, 2025, and all but RMB0.1 million of such trade receivables as of December 31, 2025 were aged within 90 days. The increase as of March 31, 2026 also reflected the rapid growth of revenue in the first quarter of 2026, certain trade receivables that remained within the credit period, and increases in other receivables and deposits relating to computing resource arrangements under which certain resources were procured or leased on a prepayment or deposit basis due to market supply conditions.",
    )

    replace_start(
        doc,
        "Our trade and other payables primarily consisted",
        "Our trade and other payables primarily consisted of trade payables, other payables, subscription received for unissued shares with preference rights, accrued staff costs, value-added tax and other tax payables and notes payables. According to the Accountants' Report, our trade and other payables increased from RMB8.2 million as of December 31, 2023 to RMB47.5 million as of December 31, 2024 and RMB238.4 million as of December 31, 2025. Trade payables increased to RMB8.8 million as of December 31, 2025 and were aged within 90 days. The Accountants' Report states that the average credit period on purchases of goods and services is 30 to 90 days. The following table sets forth a breakdown of our trade and other payables as of the dates indicated.",
    )

    replace_start(
        doc,
        "Our contract liabilities represented",
        "Our contract liabilities represented advances received from customers. Our contract liabilities amounted to nil, RMB1.1 million, RMB13.0 million and RMB38.1 million as of December 31, 2023, 2024, 2025 and March 31, 2026, respectively. According to the Accountants' Report, RMB1.1 million of revenue recognised in 2025 was included in contract liabilities at the beginning of the year. The increase in contract liabilities was generally in line with our business growth and reflected customer prepayments for platform software and solution projects.",
    )

    liq_anchor = replace_start(
        doc,
        "During the Track Record Period, our primary use of cash",
        "During the Track Record Period, our primary use of cash was to fund our R&D activities, employee benefit expenses, procurement and leasing of computing resources, sales and business development activities, lease payments, capital expenditures and other operational needs. We financed our operations and other capital requirements mainly through equity financing, bank borrowings and, to a lesser extent, cash generated from revenue-generating activities. According to the Accountants' Report, our cash and cash equivalents were short-term bank deposits with an original maturity of three months or less. We also had restricted bank deposits of RMB2.2 million as of December 31, 2025, primarily restricted bank balances for the issue of bills, which carried a fixed interest rate of 1.35% per annum. We will continue to manage our cash outflows with reference to our commercialisation progress, computing resource requirements and R&D plan.",
    )
    insert_after(
        liq_anchor,
        "According to the Accountants' Report, our deferred income represented government grants received in relation to certain government-sponsored R&D projects for front-end technologies. These grants are recorded as deferred income before completion and acceptance by the government of the related R&D projects, and are transferred to income in the form of reduced depreciation charges over the useful lives of the relevant assets. Deferred income increased from RMB5.0 million as of December 31, 2023 to RMB3.9 million as of December 31, 2024 and RMB22.4 million as of December 31, 2025, primarily reflecting government grants received and amounts credited to profit or loss during the relevant years.",
    )

    replace_start(
        doc,
        "Our bank borrowings",
        "Our bank borrowings increased from nil as of December 31, 2023 to RMB32.4 million as of December 31, 2024 and further to RMB95.1 million as of December 31, 2025. According to the Accountants' Report, as of December 31, 2025, our bank borrowings consisted of RMB29.9 million of unsecured and guaranteed bank borrowings and RMB65.2 million of unsecured and unguaranteed bank borrowings, all of which were repayable within one year. Our fixed-rate bank borrowings carried effective interest rates of 2.40% in 2024 and 2.20% to 2.40% in 2025, and our variable-rate bank borrowings carried effective interest rates of 2.25% to 2.35% in 2024 and 2.18% to 2.35% in 2025, with the variable-rate borrowings based on the one-year Loan Prime Rate minus 75 and 82 basis points and reset every three months.",
    )

    replace_start(
        doc,
        "Financial Liabilities on Shares with Preferential Rights",
        "Financial Liabilities on Shares with Preferential Rights",
    )
    pref = find(doc, "Our financial liabilities on shares with preferential rights")
    mark_para(
        pref,
        "Our financial liabilities on shares with preferential rights arose from several rounds of financing through issuing shares with certain preferred rights, mainly including redemption rights, anti-dilution rights and liquidation preference rights. According to the Accountants' Report, these financing rounds included series angel, series Pre-A, series A1, series A2, strategic series and series A+. The redemption rights granted to investors constitute obligations to repurchase our own equity instruments and were recognised as redemption liabilities, initially measured at fair value and subsequently measured at amortised cost. The balance increased from RMB415.5 million as of December 31, 2023 to RMB932.0 million as of December 31, 2024 and RMB1,608.6 million as of December 31, 2025, primarily due to additions of RMB452.4 million and RMB565.3 million in 2024 and 2025, respectively, and finance costs charged on the liabilities of RMB64.1 million and RMB111.3 million in 2024 and 2025, respectively. [Company to confirm status of termination or conversion of preferential rights upon Listing.]",
    )

    update_table_cells(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
