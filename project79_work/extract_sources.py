from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from pypdf import PdfReader


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def text_from_el(el: ET.Element) -> str:
    return "".join(t.text or "" for t in el.findall(".//w:t", NS))


def extract_docx(path: Path) -> dict:
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)

    comments = []
    revisions = []
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        if "word/comments.xml" in names:
            root = ET.fromstring(zf.read("word/comments.xml"))
            for c in root.findall(".//w:comment", NS):
                comments.append(
                    {
                        "id": c.attrib.get(f"{{{NS['w']}}}id"),
                        "author": c.attrib.get(f"{{{NS['w']}}}author"),
                        "date": c.attrib.get(f"{{{NS['w']}}}date"),
                        "text": text_from_el(c),
                    }
                )
        if "word/document.xml" in names:
            root = ET.fromstring(zf.read("word/document.xml"))
            for tag in ("ins", "del"):
                for el in root.findall(f".//w:{tag}", NS):
                    t = text_from_el(el)
                    if t.strip():
                        revisions.append(
                            {
                                "type": tag,
                                "author": el.attrib.get(f"{{{NS['w']}}}author"),
                                "date": el.attrib.get(f"{{{NS['w']}}}date"),
                                "text": t,
                            }
                        )
    return {"paragraphs": paras, "tables": tables, "comments": comments, "revisions": revisions}


def extract_pdf(path: Path) -> list[str]:
    reader = PdfReader(str(path))
    pages = []
    for i, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        text = re.sub(r"[ \t]+", " ", text)
        pages.append(f"--- Page {i} ---\n{text.strip()}")
    return pages


def write_text(path: Path, data: dict | list[str]) -> None:
    if isinstance(data, list):
        path.write_text("\n\n".join(data), encoding="utf-8")
    else:
        chunks = []
        chunks.append("# Paragraphs\n")
        for i, p in enumerate(data["paragraphs"], 1):
            chunks.append(f"[P{i}] {p}")
        chunks.append("\n# Tables\n")
        for ti, table in enumerate(data["tables"], 1):
            chunks.append(f"\n[T{ti}]")
            for row in table:
                chunks.append(" | ".join(row))
        chunks.append("\n# Comments\n")
        for c in data["comments"]:
            chunks.append(json.dumps(c, ensure_ascii=False))
        chunks.append("\n# Revisions\n")
        for r in data["revisions"]:
            chunks.append(json.dumps(r, ensure_ascii=False))
        path.write_text("\n".join(chunks), encoding="utf-8")


def main() -> None:
    base = Path("/Users/shixiaolei/Documents/Circleup/project79_work")
    sources = {
        "financial": Path("/Users/shixiaolei/Desktop/Project 79 - Financial Information (KE draft 202605xx) - 0521_YL.docx"),
        "fddq": Path("/Users/shixiaolei/Desktop/79项目  财务尽职调查问卷 (V1.0)-访谈纪要-20260527整合.docx"),
        "pdf": Path("/Users/shixiaolei/Downloads/EN Full Set_@HIP25080008_E_Galaxy_S_024(2232)_cs.pdf"),
    }
    write_text(base / "financial_extracted.txt", extract_docx(sources["financial"]))
    write_text(base / "fddq_extracted.txt", extract_docx(sources["fddq"]))
    write_text(base / "pdf_extracted.txt", extract_pdf(sources["pdf"]))


if __name__ == "__main__":
    main()
