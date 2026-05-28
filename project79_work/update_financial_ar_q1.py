from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

BLUE = RGBColor(0x00, 0x2F, 0x6C)

SRC = Path("/Users/shixiaolei/Documents/Circleup/project79_work/Project 79 - Financial Information revised markup 20260528.docx")
OUT = Path("/Users/shixiaolei/Documents/Circleup/project79_work/Project 79 - Financial Information revised markup AR Q1 20260528.docx")


def mark(p, text):
    p.clear()
    r = p.add_run(text)
    r.font.color.rgb = BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def find_start(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise RuntimeError(f"not found: {prefix}")


def replace_start(doc, prefix, text):
    p = find_start(doc, prefix)
    mark(p, text)
    return p


def insert_after(p, text):
    from copy import deepcopy
    from docx.text.paragraph import Paragraph

    new_p = deepcopy(p._p)
    p._p.addnext(new_p)
    np = Paragraph(new_p, p._parent)
    np.clear()
    mark(np, text)
    return np


def color_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = BLUE


def set_row_values(row, values):
    for i, value in enumerate(values):
        if i >= len(row.cells):
            break
        row.cells[i].text = value
        color_cell(row.cells[i])


def update_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            label = cells[0]
            if label == "Deferred income" and len(cells) >= 6:
                # Discussion of major BS current liabilities table.
                set_row_values(row, ["Deferred income", "—", "—", "—", "918", "—", "—"][: len(row.cells)])
            elif label == "Total current liabilities" and len(cells) >= 6:
                # Use latest Q1 main table presentation, which reclassifies current deferred income.
                set_row_values(row, ["Total current liabilities", "11,238", "11,238", "11,238", "189,640", "442,914", "571,918"][: len(row.cells)])
            elif label == "Net current assets" and len(cells) >= 6:
                set_row_values(row, ["Net current assets", "246,842", "246,842", "246,842", "249,975", "407,118", "334,346"][: len(row.cells)])
            elif label == "Total assets less current liabilities" and len(cells) >= 6:
                set_row_values(row, ["Total assets less current liabilities", "396,099", "396,099", "396,099", "972,279", "1,039,305", "979,447"][: len(row.cells)])
            elif label == "Deferred income" and len(cells) == 5:
                set_row_values(row, ["Deferred income", "5,000", "5,000", "3,005", "22,393", "55,354"])
            elif label == "Total" and len(cells) == 5 and any("Financial liabilities on shares" in c.text for r in table.rows for c in r.cells):
                # Indebtedness total table.
                if cells[-1] == "2,233,418":
                    set_row_values(row, ["Total", "421,379", "1,481,172", "1,481,172", "2,136,120", "2,233,418"])
            elif label == "Gross profit margin(1)" and len(cells) >= 5:
                set_row_values(row, ["Gross profit margin(1)", "—", "77.7%", "55.2%", "[●]%", "20.6%"])


def main():
    doc = Document(SRC)

    replace_start(
        doc,
        "The revenue from our agentic platform software is recognized",
        "Revenue from our agentic platform software is recognised at a point in time when the software is delivered to the customer, inspected and accepted by the customer. According to the Accountants' Report, payment is generally due within one to three months from delivery.",
    )
    replace_start(
        doc,
        "The revenue from our AGI enterprise solutions is recognized",
        "Revenue from our AGI enterprise solutions is recognised at a point in time when the software platform and related services are delivered to the customer's designated place, inspected and accepted by the customer. According to the Accountants' Report, payment is generally due within one to three months from delivery.",
    )
    replace_start(
        doc,
        "The revenue from our Agentic MaaS is recognized",
        "Revenue from our Agentic MaaS is recognised as services are rendered based on the customer's consumption of resources for usage-based contracts.",
    )

    replace_start(
        doc,
        "Our other income primarily consisted",
        "Our other income primarily consisted of (i) government grants, (ii) interest income and (iii) others. According to the Accountants' Report, our government grants amounted to RMB1.2 million, RMB22.0 million and RMB23.2 million in 2023, 2024 and 2025, respectively, and mainly represented subsidies received from local governments in Beijing and Shanghai to reward contributions by our subsidiaries to local economic growth or for industry development. The Accountants' Report states that there were no unfulfilled conditions or contingencies relating to these grants. Our interest income amounted to nil, RMB7.5 million and RMB8.3 million in 2023, 2024 and 2025, respectively. Others amounted to nil, RMB7.8 million and RMB4.2 million in 2023, 2024 and 2025, respectively. Other income amounted to RMB9.0 million for the three months ended March 31, 2026 according to the 2026 Q1 main statements. The following table sets forth the breakdown of our other income in absolute amount and as a percentage of total other income for the periods indicated.",
    )

    replace_start(
        doc,
        "Our finance costs primarily represented",
        "Our finance costs primarily represented interest on financial liabilities on shares with preferential rights, lease liabilities and bank borrowings. We incurred finance costs of RMB9.9 million, RMB70.3 million, RMB134.1 million, RMB[●] million and RMB44.9 million in 2023, 2024 and 2025 and the three months ended March 31, 2025 and 2026, respectively. According to the Accountants' Report, interest on financial liabilities on shares with preferential rights was RMB9.9 million, RMB64.1 million and RMB111.3 million in 2023, 2024 and 2025, respectively; interest on lease liabilities was nil, RMB6.3 million and RMB21.2 million in 2023, 2024 and 2025, respectively; and interest on bank borrowings was nil, RMB25 thousand and RMB1.6 million in 2023, 2024 and 2025, respectively. According to the 2026 Q1 main statements, our finance costs were RMB44.9 million for the three months ended March 31, 2026. The following table sets forth a breakdown of our finance costs for the periods indicated.",
    )

    replace_start(
        doc,
        "Under the Law of the PRC on Enterprise Income Tax",
        "Under the Law of the PRC on Enterprise Income Tax (\"EIT Law\") and Implementation Regulation of the EIT Law, the tax rate of the PRC subsidiaries is generally 25% during the Track Record Period, except for entities entitled to preferential tax treatment. According to the Accountants' Report, Shanghai Infinigence was qualified as a high and new technology enterprise (\"HNTE\") on December 25, 2025 and Beijing Wuwen was qualified as an HNTE on December 30, 2025, and each was entitled to a preferential tax rate of 15% from 2025 to 2027. Beijing Wuwen and Shanghai Infinigence were also entitled to the R&D expense super deduction during the Track Record Period.",
    )

    replace_start(
        doc,
        "Our income tax credit increased by 49.1%",
        "Our income tax credit increased by 49.1% from RMB2.2 million in 2024 to RMB3.4 million in 2025, primarily due to movements in deferred tax arising from right-of-use assets, lease liabilities, share-based payments and allowances for credit losses, as set out in the Accountants' Report.",
    )
    replace_start(
        doc,
        "Our income tax credit decreased by 48.4%",
        "Our income tax credit decreased by 48.4% from RMB4.4 million in 2023 to RMB2.2 million in 2024, primarily due to movements in deferred tax arising from right-of-use assets, lease liabilities, share-based payments and allowances for credit losses, as set out in the Accountants' Report.",
    )

    replace_start(
        doc,
        "Our property, plant and equipment primarily consisted",
        "Our property, plant and equipment primarily consisted of electronic equipment, office furniture and others. Our electronic equipment mainly included GPU servers, CPU servers, storage equipment, network security equipment and other equipment used to support our R&D activities and internal computing environment. According to the Accountants' Report, our electronic equipment is depreciated at 33.33% per annum, office furniture at 20% to 25% per annum and others at 33.33% per annum on a straight-line basis after taking into account residual values. The following table sets forth the components of our property, plant and equipment as of the dates indicated.",
    )
    replace_start(
        doc,
        "Our right-of-use assets consisted",
        "Our right-of-use assets consisted of leased properties and server devices. According to the Accountants' Report, our right-of-use assets increased from RMB5.6 million as of December 31, 2023 to RMB538.6 million as of December 31, 2024, primarily reflecting additions of RMB591.4 million in server devices in 2024, and decreased to RMB446.7 million as of December 31, 2025, primarily reflecting depreciation charges, including RMB146.5 million for server devices in 2025, partially offset by additions. According to the 2026 Q1 main statements, our right-of-use assets increased to RMB472.3 million as of March 31, 2026. Lease contracts are generally entered into for fixed terms of one to five years and may contain extension and termination options. We generally use leasing arrangements for certain server and computing resources because leasing allows us to spread cash outflows over time and provides greater flexibility than direct purchase of relevant assets.",
    )
    replace_start(
        doc,
        "Our intangible assets consisted",
        "Our intangible assets consisted of software and patents. According to the Accountants' Report, our intangible assets increased from RMB0.01 million as of December 31, 2023 to RMB0.1 million as of December 31, 2024, and further to RMB115.4 million as of December 31, 2025, primarily due to additions of RMB12.5 million in software and RMB110.6 million in patents in 2025, partially offset by amortisation. The Accountants' Report states that these intangible assets have finite useful lives and are amortised on a straight-line basis over two to ten years. According to the 2026 Q1 main statements, our intangible assets decreased to RMB111.2 million as of March 31, 2026, primarily due to amortisation. As advised by the Company, the patent additions were related to intellectual property rights contributed by Tsinghua University and Shanghai Jiao Tong University and are related to our core R&D technologies.",
    )

    replace_start(
        doc,
        "Our inventories [increased]",
        "Our inventories amounted to RMB2.0 million as of December 31, 2025 and RMB3.1 million as of March 31, 2026 according to the 2026 Q1 main statements. Based on management's FDDQ response, inventories primarily represented costs incurred for solution projects before the relevant revenue recognition milestone was reached and are expected to be transferred to cost of sales when the relevant project reaches the revenue recognition point. No inventory balance was separately presented in the Accountants' Report for 2023 to 2025, and the relevant prior-period balances should remain blank unless and until confirmed by the Reporting Accountants.",
    )

    replace_start(
        doc,
        "Our trade and other receivables increased",
        "Our trade and other receivables increased from RMB50.3 million as of December 31, 2023 to RMB83.9 million as of December 31, 2024, and further to RMB187.2 million and RMB337.8 million as of December 31, 2025 and March 31, 2026, respectively, primarily in line with our business expansion and revenue growth. According to the Accountants' Report, trade receivables net of allowances increased from nil as of December 31, 2023 to RMB0.6 million as of December 31, 2024 and RMB41.4 million as of December 31, 2025, and all but RMB0.1 million of such trade receivables as of December 31, 2025 were aged within 90 days. The increase as of March 31, 2026 also reflected the rapid growth of revenue in the first quarter of 2026, certain trade receivables that remained within the credit period, and increases in other receivables and deposits relating to computing resource arrangements under which certain resources were procured or leased on a prepayment or deposit basis due to market supply conditions.",
    )
    replace_start(
        doc,
        "Our trade and other payables primarily consisted",
        "Our trade and other payables primarily consisted of trade payables, other payables, subscription received for unissued shares with preference rights, accrued staff costs, value-added tax and other tax payables and notes payables. According to the Accountants' Report, our trade and other payables increased from RMB8.2 million as of December 31, 2023 to RMB47.5 million as of December 31, 2024 and RMB238.4 million as of December 31, 2025. According to the 2026 Q1 main statements, our trade and other payables further increased to RMB295.2 million as of March 31, 2026. Trade payables increased to RMB8.8 million as of December 31, 2025 and were aged within 90 days. The Accountants' Report states that the average credit period on purchases of goods and services is 30 to 90 days. The following table sets forth a breakdown of our trade and other payables as of the dates indicated.",
    )
    replace_start(
        doc,
        "Our trade and other payables increased",
        "Our trade and other payables increased from RMB8.2 million as of December 31, 2023 to RMB47.5 million as of December 31, 2024, further to RMB238.4 million and RMB295.2 million as of December 31, 2025 and March 31, 2026, respectively, primarily in line with our business expansion, increased procurement of computing resources and daily operating purchases, and the increase in accrued staff costs and other payables.",
    )
    replace_start(
        doc,
        "Our contract liabilities represented",
        "Our contract liabilities represented advances received from customers. Our contract liabilities amounted to nil, RMB1.1 million, RMB13.0 million and RMB38.1 million as of December 31, 2023, 2024, 2025 and March 31, 2026, respectively. According to the Accountants' Report, RMB1.1 million of revenue recognised in 2025 was included in contract liabilities at the beginning of the year. The increase in contract liabilities was generally in line with our business growth and reflected customer prepayments for platform software and solution projects.",
    )

    liq = replace_start(
        doc,
        "During the Track Record Period, our primary use of cash",
        "During the Track Record Period, our primary use of cash was to fund our R&D activities, employee benefit expenses, procurement and leasing of computing resources, sales and business development activities, lease payments, capital expenditures and other operational needs. We financed our operations and other capital requirements mainly through equity financing, bank borrowings and, to a lesser extent, cash generated from revenue-generating activities. According to the Accountants' Report, our cash and cash equivalents were short-term bank deposits with an original maturity of three months or less. We also had restricted bank deposits of RMB2.2 million as of December 31, 2025 and March 31, 2026, primarily restricted bank balances for the issue of bills, which carried a fixed interest rate of 1.35% per annum as of December 31, 2025. We will continue to manage our cash outflows with reference to our commercialisation progress, computing resource requirements and R&D plan.",
    )
    insert_after(
        liq,
        "According to the Accountants' Report, our deferred income represented government grants received in relation to certain government-sponsored R&D projects for front-end technologies. These grants are recorded as deferred income before completion and acceptance by the government of the related R&D projects, and are transferred to income in the form of reduced depreciation charges over the useful lives of the relevant assets. Deferred income was RMB5.0 million, RMB3.9 million and RMB22.4 million as of December 31, 2023, 2024 and 2025, respectively. According to the 2026 Q1 main statements, non-current deferred income increased to RMB55.4 million as of March 31, 2026.",
    )

    replace_start(
        doc,
        "Our net current assets decreased",
        "Our net current assets decreased from RMB407.1 million as of December 31, 2025 to RMB334.3 million as of March 31, 2026, primarily due to a decrease in term deposits of RMB58.5 million, an increase in trade and other payables of RMB56.9 million, an increase in contract liabilities of RMB25.2 million, an increase in current lease liabilities of RMB25.8 million and a decrease in cash and cash equivalents of RMB37.0 million, partially offset by an increase in trade and other receivables of RMB150.6 million.",
    )
    replace_start(
        doc,
        "Our net current assets increased",
        "Our net current assets increased from RMB250.0 million as of December 31, 2024 to RMB407.1 million as of December 31, 2025, primarily due to an increase in cash and cash equivalents of RMB211.3 million and an increase in trade and other receivables of RMB103.3 million, partially offset by an increase in trade and other payables of RMB190.9 million, an increase in bank borrowings of RMB62.8 million and an increase in contract liabilities of RMB11.9 million.",
    )

    replace_start(
        doc,
        "Our bank borrowings increased",
        "Our bank borrowings increased from nil as of December 31, 2023 to RMB32.4 million as of December 31, 2024, RMB95.1 million as of December 31, 2025 and RMB116.3 million as of March 31, 2026. According to the Accountants' Report, as of December 31, 2025, our bank borrowings consisted of RMB29.9 million of unsecured and guaranteed bank borrowings and RMB65.2 million of unsecured and unguaranteed bank borrowings, all of which were repayable within one year. Our fixed-rate bank borrowings carried effective interest rates of 2.40% in 2024 and 2.20% to 2.40% in 2025, and our variable-rate bank borrowings carried effective interest rates of 2.25% to 2.35% in 2024 and 2.18% to 2.35% in 2025, with the variable-rate borrowings based on the one-year Loan Prime Rate minus 75 and 82 basis points and reset every three months. The increase as of March 31, 2026 was primarily due to additional drawdown of bank facilities to support business operations and payment for computing resource and equipment-related needs.",
    )
    replace_start(
        doc,
        "We recorded financial liabilities on shares",
        "We recorded financial liabilities on shares with preferential rights of RMB415.5 million, RMB932.0 million, RMB1,608.6 million and RMB1,648.3 million as of December 31, 2023, 2024, 2025 and March 31, 2026, respectively. According to the Accountants' Report, these liabilities arose from several rounds of financing through issuing shares with certain preferred rights, mainly including redemption rights, anti-dilution rights and liquidation preference rights. The redemption rights granted to investors constitute obligations to repurchase our own equity instruments and were recognised as redemption liabilities, initially measured at fair value and subsequently measured at amortised cost. The increase from 2023 to 2025 was primarily due to additions of RMB452.4 million and RMB565.3 million in 2024 and 2025, respectively, and finance costs charged on the liabilities of RMB64.1 million and RMB111.3 million in 2024 and 2025, respectively. The increase in the first quarter of 2026 was primarily due to finance costs accrued over time. [Company to confirm status of termination or conversion of preferential rights upon Listing.]",
    )

    update_tables(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
